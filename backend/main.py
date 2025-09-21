"""
FastAPI backend for GenAI Assistant
- Hybrid OCR (EasyOCR + PaddleOCR) with graceful fallbacks
- RAG (local embeddings + optional FAISS)
- Agentic planning (optional)
- Report PDF builder
Reads OPENAI_API_KEY from .env at project root.
"""
import os, io, zipfile, tempfile, warnings, math, json, uuid, base64
from typing import List, Dict, Any, Optional, Tuple
from types import SimpleNamespace

# --- load .env
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

import numpy as np
import pandas as pd
from sklearn.metrics.pairwise import cosine_similarity
from PIL import Image
import pdfplumber, fitz 
from docx import Document

from fastapi import FastAPI, UploadFile, File, Form, Body
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware

# OpenAI (Embeddings + Chat) 
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
LLM_AVAILABLE = bool(OPENAI_API_KEY)

from openai import OpenAI
client = OpenAI(api_key=OPENAI_API_KEY or None)

# Autogen AgentChat (agentic flow)
AGENTIC_OK = True
try:
    from autogen_agentchat.agents import AssistantAgent
    from autogen_ext.models.openai import OpenAIChatCompletionClient
    from autogen_agentchat.teams import RoundRobinGroupChat
    from autogen_agentchat.conditions import MaxMessageTermination
except Exception:
    AGENTIC_OK = False

if not LLM_AVAILABLE:
    AGENTIC_OK = False

# Optional FAISS 
try:
    import faiss
    FAISS_OK = True
except Exception:
    FAISS_OK = False

# Suppress & patch torch.load warning 
warnings.filterwarnings("ignore", category=FutureWarning, message=r".*torch\.load.*weights_only.*")
EASY_TORCH_PATCH_ACTIVE = False
try:
    import torch
    _ORIG_TORCH_LOAD = torch.load
    def _patched_torch_load(*args, **kwargs):
        if "weights_only" not in kwargs:
            kwargs["weights_only"] = True
        return _ORIG_TORCH_LOAD(*args, **kwargs)
    torch.load = _patched_torch_load  
    EASY_TORCH_PATCH_ACTIVE = True
except Exception:
    pass

# OCR backends (EasyOCR + PaddleOCR)
OCR_AVAILABLE = False
OCR_REASON: List[str] = []
EASYOCR_READER = None
PADDLE_AVAILABLE = False
PADDLE_OCR = None
EASY_DEVICE = "CPU"

try:
    import cv2
except Exception as e:
    cv2 = None
    OCR_REASON.append(f"opencv import failed: {e}")

try:
    import easyocr  
    try:
        use_gpu = ('torch' in globals()) and torch.cuda.is_available()
        EASY_DEVICE = "GPU" if use_gpu else "CPU"
    except Exception:
        use_gpu = False
        EASY_DEVICE = "CPU"
    EASYOCR_READER = easyocr.Reader(['en'], gpu=use_gpu)
except Exception as e:
    OCR_REASON.append(f"easyocr init failed: {e}")
    EASYOCR_READER = None

try:
    from paddleocr import PaddleOCR  
    PADDLE_OCR = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)
    PADDLE_AVAILABLE = True
except Exception as e:
    OCR_REASON.append(f"paddleocr unavailable: {e}")
    PADDLE_OCR = None
    PADDLE_AVAILABLE = False

OCR_AVAILABLE = (EASYOCR_READER is not None) or PADDLE_AVAILABLE

# Report export
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table as RLTable, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

# Embeddings 
def embed_texts(texts: List[str]) -> np.ndarray:
    try:
        if not texts:
            return np.zeros((0, 1536), dtype=np.float32)
        resp = client.embeddings.create(model="text-embedding-3-small", input=texts)
        return np.array([d.embedding for d in resp.data], dtype=np.float32)
    except Exception:
        return np.zeros((len(texts), 1536), dtype=np.float32)

# Chunking 
def chunk_text(text: str, max_tokens: int = 600, overlap: int = 80) -> List[str]:
    if not text:
        return []
    chunk_size = max(200, max_tokens * 4)
    stride = max(50, chunk_size - overlap * 4)
    out, i, n = [], 0, len(text)
    while i < n:
        out.append(text[i:i + chunk_size])
        i += stride
    return out

# OCR helpers
def _is_table_like(gray: np.ndarray) -> bool:
    if cv2 is None or gray is None:
        return False
    edges = cv2.Canny(gray, 80, 160)
    lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=120, minLineLength=40, maxLineGap=5)
    if lines is None:
        return False
    hlines, vlines = 0, 0
    for l in lines[:, 0, :]:
        x1, y1, x2, y2 = l
        import math as _m
        angle = _m.degrees(_m.atan2(y2 - y1, x2 - x1))
        if abs(angle) < 10:
            hlines += 1
        elif abs(abs(angle) - 90) < 10:
            vlines += 1
    return (hlines + vlines) >= 10

def _preprocess_common(img_bytes: bytes):
    if cv2 is None:
        return None, None
    import numpy as _np
    arr = _np.frombuffer(img_bytes, _np.uint8)
    bgr = cv2.imdecode(arr, cv2.IMREAD_COLOR)
    if bgr is None:
        pil = Image.open(io.BytesIO(img_bytes)).convert("RGB")
        bgr = cv2.cvtColor(np.array(pil), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(bgr, cv2.COLOR_BGR2GRAY)
    h, w = gray.shape[:2]
    scale = 1.6 if max(h, w) < 1500 else 1.0
    if scale != 1.0:
        bgr = cv2.resize(bgr, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_LINEAR)
        gray = cv2.resize(gray, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_LINEAR)
    gray = cv2.fastNlMeansDenoising(gray, None, 15, 7, 21)
    return bgr, gray

def _preprocess_for_easyocr(gray: np.ndarray) -> np.ndarray:
    if cv2 is None:
        return gray
    return cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
                                 cv2.THRESH_BINARY, 31, 10)

def _clean_numeric_tokens(text: str) -> str:
    repl = (("O", "0"), ("o", "0"), ("l", "1"), ("I", "1"), ("S", "5"))
    out = []
    for tok in text.split():
        t = tok
        if any(c.isdigit() for c in t) or "%" in t:
            for a, b in repl:
                t = t.replace(a, b)
            t = t.strip(",;:。．٫")
        out.append(t)
    return " ".join(out)

def easyocr_text_from_bytes(img_bytes: bytes):
    if EASYOCR_READER is None or cv2 is None:
        return "", ""
    try:
        _, gray = _preprocess_common(img_bytes)
        if gray is None:
            return "", ""
        proc = _preprocess_for_easyocr(gray)
        raw_list = EASYOCR_READER.readtext(proc, detail=0, paragraph=True)
        raw = " ".join(raw_list).strip()
        return _clean_numeric_tokens(raw), raw
    except Exception:
        try:
            pil = Image.open(io.BytesIO(img_bytes)).convert("RGB")
            raw_list = EASYOCR_READER.readtext(np.array(pil), detail=0, paragraph=True)
            raw = " ".join(raw_list).strip()
            return _clean_numeric_tokens(raw), raw
        except Exception:
            return "", ""

def paddle_text_and_cells_from_bytes(img_bytes: bytes):
    if not PADDLE_AVAILABLE:
        return "", []
    try:
        bgr, gray = _preprocess_common(img_bytes)
        if bgr is None:
            return "", []
        result = PADDLE_OCR.ocr(bgr, cls=True)
        if not result or not result[0]:
            return "", []
        lines, texts = [], []
        for box, (txt, conf) in result[0]:
            if not txt:
                continue
            xs = [p[0] for p in box]; ys = [p[1] for p in box]
            x1, x2 = min(xs), max(xs); y1, y2 = min(ys), max(ys)
            cx, cy, w, h = (x1+x2)/2.0, (y1+y2)/2.0, (x2-x1), (y2-y1)
            lines.append((txt, (cx, cy, w, h))); texts.append(txt)
        return " ".join(texts).strip(), lines
    except Exception:
        return "", []

def lines_to_dataframe(lines):
    if not lines:
        return None
    lines_sorted = sorted(lines, key=lambda t: (t[1][1], t[1][0]))
    rows = []
    tol = 14.0
    for txt, (cx, cy, w, h) in lines_sorted:
        placed = False
        for r in rows:
            _, (cx0, cy0, w0, h0) = r[-1]
            if abs(cy - cy0) <= max(tol, 0.35 * h0):
                r.append((txt, (cx, cy, w, h)))
                placed = True
                break
        if not placed:
            rows.append([(txt, (cx, cy, w, h))])
    for r in rows:
        r.sort(key=lambda t: t[1][0])
    avg_cells = sum(len(r) for r in rows) / max(1, len(rows))
    if avg_cells < 3:
        return None
    max_cols = max(len(r) for r in rows)
    data = []
    for r in rows:
        data.append([cell[0] for cell in r] + [""] * (max_cols - len(r)))
    header_candidates = data[0]
    df = pd.DataFrame(data[1:], columns=header_candidates)
    return df

def ocr_image_bytes(img_bytes: bytes):
    debug = []
    tables = []
    text_easy_clean, text_easy_raw = "", ""
    text_paddle, paddle_lines = "", []
    gray = None
    if cv2 is not None:
        try:
            _, gray = _preprocess_common(img_bytes)
        except Exception:
            gray = None
    table_like = _is_table_like(gray) if gray is not None else False
    if PADDLE_AVAILABLE and table_like:
        text_paddle, paddle_lines = paddle_text_and_cells_from_bytes(img_bytes)
        if text_paddle:
            debug.append(("Image OCR (Paddle)", "", text_paddle))
            df = lines_to_dataframe(paddle_lines)
            if df is not None and len(df) > 0:
                tables.append(df)
    if EASYOCR_READER is not None:
        text_easy_clean, text_easy_raw = easyocr_text_from_bytes(img_bytes)
        if text_easy_clean or text_easy_raw:
            debug.append(("Image OCR (EasyOCR)", text_easy_raw, text_easy_clean))
    if PADDLE_AVAILABLE and not table_like and not text_paddle:
        text_paddle, paddle_lines = paddle_text_and_cells_from_bytes(img_bytes)
        if text_paddle:
            debug.append(("Image OCR (Paddle)", "", text_paddle))
            df = lines_to_dataframe(paddle_lines)
            if df is not None and len(df) > 0:
                tables.append(df)
    merged_clean = " ".join(s for s in [text_easy_clean, text_paddle] if s).strip()
    merged_raw = " ".join(s for s in [text_easy_raw, text_paddle] if s).strip()
    return merged_clean, merged_raw, tables, debug

# Loaders 
def load_pdf_text_tables_images(path: str, enable_pdf_ocr: bool = True):
    text_parts, tables, images, ocr_debug = [], [], [], []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
                elif enable_pdf_ocr and OCR_AVAILABLE:
                    pil_img = page.to_image(resolution=220).original
                    buf = io.BytesIO(); pil_img.save(buf, "PNG")
                    page_png = buf.getvalue()
                    images.append(page_png)
                    cleaned, raw, tbs, dbg = ocr_image_bytes(page_png)
                    if cleaned:
                        text_parts.append(cleaned)
                    if tbs:
                        tables.extend(tbs)
                    for d in dbg:
                        ocr_debug.append(("PDF page " + d[0], d[1], d[2]))
                try:
                    for t in (page.extract_tables() or []):
                        if t and any(any(cell for cell in row) for row in t):
                            df = pd.DataFrame(t[1:], columns=t[0] if t[0] else None)
                            tables.append(df)
                except Exception:
                    pass
    except Exception:
        pass
    try:
        doc = fitz.open(path)
        for p in range(len(doc)):
            for img in doc[p].get_images(full=True):
                xref = img[0]
                base_img = doc.extract_image(xref)
                img_bytes = base_img.get("image", b"")
                if img_bytes:
                    images.append(img_bytes)
                    if OCR_AVAILABLE:
                        cleaned, raw, tbs, dbg = ocr_image_bytes(img_bytes)
                        if cleaned:
                            text_parts.append(cleaned)
                        if tbs:
                            tables.extend(tbs)
                        for d in dbg:
                            ocr_debug.append(("PDF embed " + d[0], d[1], d[2]))
    except Exception:
        pass
    return "\n".join(text_parts), tables, images, ocr_debug

def load_docx(path: str, enable_docx_img_ocr: bool = True):
    text_parts, tables, images, ocr_debug = [], [], [], []
    try:
        doc = Document(path)
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)
        for tbl in doc.tables:
            data = [[c.text for c in r.cells] for r in tbl.rows]
            if data:
                df = pd.DataFrame(data[1:], columns=data[0] if data[0] else None)
                tables.append(df)
    except Exception:
        pass
    if enable_docx_img_ocr:
        try:
            with zipfile.ZipFile(path) as z:
                for name in z.namelist():
                    if name.lower().startswith("word/media/"):
                        img_bytes = z.read(name)
                        images.append(img_bytes)
                        if OCR_AVAILABLE:
                            cleaned, raw, tbs, dbg = ocr_image_bytes(img_bytes)
                            if cleaned:
                                text_parts.append(cleaned)
                            if tbs:
                                tables.extend(tbs)
                            for d in dbg:
                                ocr_debug.append(("DOCX embed " + d[0], d[1], d[2]))
        except Exception:
            pass
    return "\n".join(text_parts), tables, images, ocr_debug

def load_image(path: str, enable_img_ocr: bool = True):
    try:
        with open(path, "rb") as f:
            img_bytes = f.read()
    except Exception:
        return "", [], [], []
    text, debug = "", []
    tables: List[pd.DataFrame] = []
    if enable_img_ocr and OCR_AVAILABLE:
        cleaned, raw, tbs, dbg = ocr_image_bytes(img_bytes)
        text = cleaned
        tables = tbs
        debug.extend(dbg)
    return text, tables, [img_bytes], debug

def load_excel_csv(path: str):
    try:
        if path.lower().endswith((".xlsx", ".xls")):
            xls = pd.ExcelFile(path)
            frames = [xls.parse(s) for s in xls.sheet_names]
        else:
            frames = [pd.read_csv(path)]
        text = "\n\n".join(df.to_csv(index=False) for df in frames)
        return text, frames, [], []
    except Exception:
        return "", [], [], []

def load_any(path: str, enable_img_ocr: bool, enable_pdf_ocr: bool, enable_docx_img_ocr: bool):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return load_pdf_text_tables_images(path, enable_pdf_ocr=enable_pdf_ocr)
    if ext == ".docx":
        return load_docx(path, enable_docx_img_ocr=enable_docx_img_ocr)
    if ext in {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif"}:
        return load_image(path, enable_img_ocr=enable_img_ocr)
    if ext in {".csv", ".xlsx", ".xls"}:
        return load_excel_csv(path)
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read(), [], [], []
    except Exception:
        return "", [], [], []

# RAG 
class RAGIndex:
    def __init__(self):
        self.chunks: List[str] = []
        self.chunk_meta: List[Dict[str, Any]] = []
        self.emb: Optional[np.ndarray] = None
        self.faiss_index = None

    def add_document(self, filename: str, text: str):
        if not text:
            return
        chunks = chunk_text(text)
        if not chunks:
            return
        start = len(self.chunks)
        self.chunks.extend(chunks)
        self.chunk_meta.extend([{"filename": filename, "chunk_id": i} for i in range(start, len(self.chunks))])

    def build(self):
        if not self.chunks:
            self.emb = None
            self.faiss_index = None
            return
        vectors = embed_texts(self.chunks)
        self.emb = vectors
        if FAISS_OK and vectors.size > 0:
            dim = vectors.shape[1]
            index = faiss.IndexFlatIP(dim)
            faiss.normalize_L2(vectors)
            index.add(vectors)
            self.faiss_index = index

    def query(self, q: str, k: int = 5):
        if not self.chunks:
            return []
        qv = embed_texts([q]).astype(np.float32)
        if qv.size == 0:
            return []
        if FAISS_OK and self.faiss_index is not None:
            import numpy as _np
            qv_norm = qv / _np.linalg.norm(qv, axis=1, keepdims=True)
            D, I = self.faiss_index.search(qv_norm, k)
            return [(self.chunks[i], self.chunk_meta[i], float(D[0][j])) for j, i in enumerate(I[0]) if i >= 0]
        sims = cosine_similarity(qv, self.emb)[0]
        import numpy as _np
        idxs = _np.argsort(-sims)[:k]
        return [(self.chunks[i], self.chunk_meta[i], float(sims[i])) for i in idxs]

SYSTEM_PROMPT = (
    "A clinical/document AI assistant. Use retrieved context when helpful. "
    "Write clear, well-structured, factual prose. Prefer concise sentences. "
    "Do not copy large raw blocks; paraphrase and summarize unless the user asked for verbatim."
)

def _llm_stub(msg: str = "LLM unavailable. Provided OPENAI_API_KEY via .env.") -> Any:
    return SimpleNamespace(choices=[SimpleNamespace(message=SimpleNamespace(content=msg))])

def llm_chat(messages: List[Dict[str, str]], temperature: float = 0.2):
    if not LLM_AVAILABLE:
        return _llm_stub()
    try:
        return client.chat.completions.create(model="gpt-4o-mini", temperature=temperature, messages=messages)
    except Exception:
        try:
            return client.chat.completions.create(model="gpt-4o-mini", temperature=0.0, messages=messages)
        except Exception:
            return _llm_stub()

def df_to_csv_clip(df: pd.DataFrame, max_rows: int = 30, max_cols: int = 12) -> str:
    _df = df.copy().iloc[:max_rows, :max_cols]
    return _df.to_csv(index=False)

# Agents & Orchestrator
def build_agents():
    if not AGENTIC_OK:
        return None
    model_client = OpenAIChatCompletionClient(model="gpt-4o-mini", api_key=OPENAI_API_KEY)
    orchestrator = AssistantAgent(
        name="orchestrator",
        system_message=(
            "You are the Controller Agent. Break the user's goal into steps and coordinate agents. "
            "When asked for a plan, respond with compact JSON containing keys like: "
            "steps, per_section_prompts, sec_queries, extraction_queries, summary_prompt."
        ),
        model_client=model_client,
    )
    qa_agent = AssistantAgent(
        name="qa_agent",
        system_message=(
            "Q&A Agent (RAG). Given a user question or section names, propose sharp retrieval queries. "
            "Return JSON: { 'sec_queries': {sectionOrUserQ: [query1, query2]} } or { 'queries': [q1,q2] }"
        ),
        model_client=model_client,
    )
    extraction = AssistantAgent(
        name="extraction_agent",
        system_message=(
            "Extraction Agent. Suggest what exact text lines/snippets, table topics/headers, and image cues to pull. "
            "Return JSON: { 'extraction_queries': { section: {'text': [...], 'tables':[...], 'images':[...]} } }"
        ),
        model_client=model_client,
    )
    assembly = AssistantAgent(
        name="report_assembly",
        system_message=(
            "Report Assembly Agent. Draft concise prompts for LLM writing per section. "
            "Return JSON: { 'per_section_prompts': { section: '<prompt>' } }"
        ),
        model_client=model_client,
    )
    summarizer = AssistantAgent(
        name="summarization_agent",
        system_message=(
            "Summarization Agent. Provide a one-line instruction to create a 5–8 bullet HTML summary. "
            "Return JSON: { 'summary_prompt': '<instruction>' }"
        ),
        model_client=model_client,
    )
    team_plan = RoundRobinGroupChat([orchestrator, qa_agent, extraction, assembly, summarizer],
                                    termination_condition=MaxMessageTermination(max_messages=8))
    team_qa = RoundRobinGroupChat([orchestrator, qa_agent],
                                  termination_condition=MaxMessageTermination(max_messages=3))
    return dict(team_plan=team_plan, qa_agent=qa_agent, team_qa=team_qa)

def _safe_parse_json_blob(text: str) -> Optional[Dict[str, Any]]:
    try:
        return json.loads(text)
    except Exception:
        l, r = text.find("{"), text.rfind("}")
        if l != -1 and r != -1 and r > l:
            try:
                return json.loads(text[l:r+1])
            except Exception:
                return None
        return None

# BACKEND (FastAPI)
class APISessionState:
    def __init__(self):
        self.rag = RAGIndex()
        self.upload_cache: List[Dict[str, Any]] = []
        self.messages: List[Dict[str, str]] = []

def api_top_context_for_query(state: APISessionState, query: str, k: int = 8) -> str:
    hits = state.rag.query(query, k=k)
    return "\n\n".join([h[0] for h in hits]) if hits else ""

def api_tool_extract_tables(state: APISessionState, section_query: str, top_k: int = 3) -> List[pd.DataFrame]:
    hits: List[pd.DataFrame] = []
    for item in state.upload_cache:
        for df in item.get("tables", []):
            cols = " ".join(map(str, df.columns))
            if any(w.lower() in cols.lower() for w in section_query.split()):
                hits.append(df)
    if not hits:
        for item in state.upload_cache:
            hits.extend(item.get("tables", []))
    return hits[:top_k]

def api_tool_extract_images(state: APISessionState, section_query: str, top_k: int = 3) -> List[bytes]:
    imgs: List[bytes] = []
    for item in state.upload_cache:
        imgs.extend(item.get("images", []))
    return imgs[:top_k]

def api_llm_write_section_with_tools(state: APISessionState, title: str, user_goal: str, query_hint: str, temperature: float=0.2):
    ctx = api_top_context_for_query(state, query_hint or title, k=8)
    tbls = api_tool_extract_tables(state, query_hint, top_k=3)
    imgs = api_tool_extract_images(state, query_hint, top_k=3)
    def llm_write():
        tables_payload = [f"Table {i+1} (CSV):\n{df_to_csv_clip(df)}" for i, df in enumerate(tbls[:3])]
        user_msg = (
            f"Write the '{title}' section for a clinical/document report.\n"
            f"GOAL: {user_goal}\n\n"
            f"Use these retrieved context snippets if relevant:\n{ctx[:5000]}\n\n"
            f"Tables (CSV clips):\n" + ("\n\n".join(tables_payload) if tables_payload else "None") + "\n\n"
            f"Images/graphs present: {'Yes' if imgs else 'No'}\n\n"
            "Requirements:\n"
            "• 1 short intro line, then 3–6 bullet points or 2–4 short paragraphs.\n"
            "• Precise, no fluff. Simple HTML (<p>, <ul><li>), no markdown."
        )
        msgs = [{"role": "system", "content": SYSTEM_PROMPT}, {"role": "user", "content": user_msg}]
        return llm_chat(msgs, temperature=temperature).choices[0].message.content
    html = llm_write()
    exact_text = []  
    return html, exact_text, tbls, imgs

def build_pdf_report(sections: Dict[str, Dict[str, Any]]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Small", fontSize=9, leading=11))
    flow = []
    title = sections.get("_meta_", {}).get("title", "Structured Report")
    flow.append(Paragraph(f"<b>{title}</b>", styles["Title"]))
    flow.append(Spacer(1, 12))
    for sec_name, sec in sections.items():
        if sec_name.startswith("_"):
            continue
        flow.append(Paragraph(f"<b>{sec_name}</b>", styles["Heading2"]))
        flow.append(Spacer(1, 6))
        if sec.get("narrative_html"):
            for para in sec["narrative_html"].split("\n"):
                p = para.strip()
                if not p:
                    continue
                flow.append(Paragraph(p, styles["BodyText"]))
            flow.append(Spacer(1, 6))
        for para in sec.get("exact_text", []):
            flow.append(Paragraph(para.replace("\n", "<br/>"), styles["Small"]))
            flow.append(Spacer(1, 4))
        for df, note in sec.get("tables_with_notes", []):
            data = [list(df.columns)] + df.fillna("").astype(str).values.tolist()
            table = RLTable(data, repeatRows=1)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]))
            flow.append(table)
            if note:
                flow.append(Spacer(1, 4))
                flow.append(Paragraph(note, styles["Small"]))
            flow.append(Spacer(1, 8))
        for img_bytes, cap in sec.get("images_with_captions", []):
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                    tmp.write(img_bytes)
                    tmp_path = tmp.name
                flow.append(RLImage(tmp_path, width=4.8 * inch, height=3.2 * inch))
                if cap:
                    flow.append(Spacer(1, 4))
                    flow.append(Paragraph(cap, styles["Small"]))
                flow.append(Spacer(1, 8))
            except Exception:
                pass
        if sec.get("summary"):
            flow.append(Paragraph("<i>Summary:</i>", styles["Italic"]))
            flow.append(Paragraph(sec["summary"].replace("\n", "<br/>"), styles["BodyText"]))
            flow.append(Spacer(1, 8))
    doc.build(flow)
    pdf = buf.getvalue()
    buf.close()
    return pdf

# Agentic plan (lightweight)
def api_build_report_with_agents(state: APISessionState, chosen_sections: List[str], sec_queries: Dict[str, str], user_goal: str, temperature: float):
    if not AGENTIC_OK:
        # Fallback to simple per-section writer
        out = {}
        for sec in chosen_sections:
            html, exact_snips, tbls, imgs = api_llm_write_section_with_tools(state, sec, user_goal, sec_queries.get(sec, sec), temperature)
            out[sec] = {"narrative_html": html, "exact_text": exact_snips, "tables_with_notes": [], "images_with_captions": [], "summary": ""}
        return out
    agents = build_agents()
    team_plan = agents["team_plan"]
    kickoff = (
        "User Goal:\n" + user_goal + "\n\n"
        "Sections: " + ", ".join(chosen_sections) + "\n"
        "Return compact JSON with keys: { per_section_prompts, sec_queries, summary_prompt }"
    )
    try:
        res = team_plan.run(kickoff)
        txt = res.messages[-1].content if hasattr(res, "messages") else str(res)
        parsed = _safe_parse_json_blob(txt) or {}
        per_prompts = parsed.get("per_section_prompts", {})
        sec_queries_plan = parsed.get("sec_queries", {})
    except Exception:
        per_prompts, sec_queries_plan = {}, {}
    sections_payload = {}
    for sec in chosen_sections:
        q_hint = ", ".join(sec_queries_plan.get(sec, [sec_queries.get(sec, sec)]))
        prompt_hint = per_prompts.get(sec, user_goal)
        ctx = api_top_context_for_query(state, q_hint, k=8)
        tbls = api_tool_extract_tables(state, q_hint, top_k=3)
        imgs = api_tool_extract_images(state, q_hint, top_k=3)
        tables_payload = [f"Table {i+1} (CSV):\n{df_to_csv_clip(df)}" for i, df in enumerate(tbls[:3])]
        user_msg = (
            f"Write the '{sec}' section for a clinical/document report.\n"
            f"GOAL: {prompt_hint}\n\n"
            f"Context:\n{ctx[:5000]}\n\n"
            f"Tables:\n" + ("\n\n".join(tables_payload) if tables_payload else "None") + "\n\n"
            f"Images present: {'Yes' if imgs else 'No'}\n"
            "Output simple HTML (<p>, <ul><li>), no markdown."
        )
        msgs = [{"role": "system", "content": SYSTEM_PROMPT}, {"role": "user", "content": user_msg}]
        narrative_html = llm_chat(msgs, temperature=temperature).choices[0].message.content
        sections_payload[sec] = {"narrative_html": narrative_html, "exact_text": [], "tables_with_notes": [], "images_with_captions": [], "summary": ""}
    return sections_payload

# FastAPI app + endpoints
api = FastAPI(title="GenAI Assistant API", version="1.0.0")
api.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"],
)

SESSIONS: Dict[str, APISessionState] = {}

@api.get("/api/v1/health")
def health():
    return {"status": "ok", "llm": LLM_AVAILABLE, "agentic": AGENTIC_OK, "ocr": OCR_AVAILABLE}

@api.post("/api/v1/session")
def create_session():
    sid = str(uuid.uuid4())
    SESSIONS[sid] = APISessionState()
    return {"session_id": sid}

@api.get("/api/v1/session/{session_id}/state")
def session_state(session_id: str):
    s = SESSIONS.get(session_id)
    if not s:
        return JSONResponse(status_code=404, content={"error": "unknown session"})
    return {"chunks": len(s.rag.chunks), "uploads": len(s.upload_cache), "messages": len(s.messages)}

@api.post("/api/v1/ingest")
async def ingest(
    session_id: str = Form(...),
    use_agentic_ingest: bool = Form(False),
    ocr_img_default: bool = Form(True),
    ocr_pdf_default: bool = Form(True),
    ocr_docx_default: bool = Form(True),
    files: List[UploadFile] = File(...),
):
    s = SESSIONS.get(session_id)
    if not s:
        return JSONResponse(status_code=404, content={"error": "unknown session"})
    s.upload_cache.clear()
    s.rag = RAGIndex()

    file_infos = []
    tmp_paths = []
    for up in files:
        contents = await up.read()
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(up.filename)[1]) as tmp:
            tmp.write(contents)
            path = tmp.name
        tmp_paths.append((up.filename, path))
        file_infos.append({"name": up.filename, "ext": os.path.splitext(up.filename)[1].lower(), "size": len(contents)})

    # (Optional) future: agentic ingest plan per-file based on file_infos
    plan = {"ingest_plan": {"ocr_img": ocr_img_default, "ocr_pdf": ocr_pdf_default, "ocr_docx": ocr_docx_default}}

    for (fname, path) in tmp_paths:
        ocr_img = plan["ingest_plan"]["ocr_img"]
        ocr_pdf = plan["ingest_plan"]["ocr_pdf"]
        ocr_docx = plan["ingest_plan"]["ocr_docx"]
        text, tables, images, _ = load_any(path, ocr_img, ocr_pdf, ocr_docx)
        s.upload_cache.append({"filename": fname, "path": path, "text": text, "tables": tables, "images": images})
        if text:
            s.rag.add_document(fname, text)
    s.rag.build()
    return {"ok": True, "plan": plan, "chunks": len(s.rag.chunks), "uploads": len(s.upload_cache)}

@api.post("/api/v1/chat")
def chat(payload: Dict[str, Any] = Body(...)):
    session_id = payload.get("session_id")
    q = payload.get("message", "")
    if not session_id or not q:
        return JSONResponse(status_code=400, content={"error": "session_id and message required"})
    s = SESSIONS.get(session_id)
    if not s:
        return JSONResponse(status_code=404, content={"error": "unknown session"})
    s.messages.append({"role": "user", "content": q})

    ctx_chunks: List[str] = []
    for h, _, _ in s.rag.query(q, k=6):
        ctx_chunks.append(h)
    context_text = "\n\n".join(ctx_chunks[:12])

    msgs = [{"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": f"Q: {q}\n\nContext (RAG):\n{context_text}"}]
    ans = llm_chat(msgs).choices[0].message.content if LLM_AVAILABLE else "Provide OPENAI_API_KEY via .env."
    s.messages.append({"role": "assistant", "content": ans})
    return {"answer": ans, "tokens_hint": len(context_text)}

@api.post("/api/v1/report")
def report(payload: Dict[str, Any] = Body(...)):
    session_id = payload.get("session_id")
    if not session_id:
        return JSONResponse(status_code=400, content={"error": "session_id required"})
    s = SESSIONS.get(session_id)
    if not s:
        return JSONResponse(status_code=404, content={"error": "unknown session"})
    rpt_title = payload.get("title", "Structured Document Report")
    chosen_sections = payload.get("sections", ["Introduction", "About", "Brief", "Patient Tables", "Graphs", "Summary"])
    sec_queries = payload.get("sec_queries", {sec: sec for sec in chosen_sections})
    user_goal = payload.get("user_goal", "Create a clean, concise, clinically-styled report summarizing the key content of the uploaded files.")
    use_agentic_plan = bool(payload.get("use_agentic_plan", False))
    temperature = float(payload.get("temperature", 0.2))

    sections_payload: Dict[str, Dict[str, Any]] = {"_meta_": {"title": rpt_title}}
    if use_agentic_plan:
        built = api_build_report_with_agents(s, chosen_sections, sec_queries, user_goal, temperature)
        sections_payload.update(built)
    else:
        keep_for_summary: Dict[str, Dict[str, Any]] = {}
        for sec in chosen_sections:
            html, exact_snips, tbls, imgs = api_llm_write_section_with_tools(s, sec, user_goal, sec_queries.get(sec, sec), temperature=temperature)
            tables_with_notes = []
            if ("table" in sec.lower()) and tbls:
                for df in tbls[:3]:
                    tables_with_notes.append((df.head(50), ""))
            elif tbls[:1]:
                tables_with_notes.append((tbls[0].head(50), ""))
            images_with_caps = []
            if imgs[:1]:
                images_with_caps.append((imgs[0], ""))
            blk = {"narrative_html": html, "exact_text": exact_snips, "tables_with_notes": tables_with_notes, "images_with_captions": images_with_caps, "summary": ""}
            sections_payload[sec] = blk
            keep_for_summary[sec] = blk

        if "Summary" in chosen_sections:
            combined_for_summary = []
            for _, blk in keep_for_summary.items():
                if blk.get("narrative_html"):
                    combined_for_summary.append(blk["narrative_html"])
                combined_for_summary.extend(blk.get("exact_text", []))
            combined_text = "\n\n".join(combined_for_summary)[-8000:]
            summary_text = ""
            if combined_text:
                msgs = [{"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user", "content": "Write 5–8 bullet points as <ul><li>…</li></ul>.\n\n" + combined_text}]
                try:
                    summary_text = llm_chat(msgs, temperature=temperature).choices[0].message.content
                except Exception:
                    summary_text = ""
            if "Summary" in sections_payload:
                sections_payload["Summary"]["summary"] = summary_text
            else:
                sections_payload["Summary"] = {"narrative_html": "", "exact_text": [], "tables_with_notes": [], "images_with_captions": [], "summary": summary_text}

    pdf_bytes = build_pdf_report(sections_payload)
    return StreamingResponse(io.BytesIO(pdf_bytes), media_type="application/pdf",
                             headers={"Content-Disposition": 'attachment; filename="report.pdf"'})
